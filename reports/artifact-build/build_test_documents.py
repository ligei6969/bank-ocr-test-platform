"""Build test-plan and test-case-template DOCX artifacts for the OCR platform."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT_DIR = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT_DIR / "reports" / "test-documents"
PLAN_PATH = OUTPUT_DIR / "银行OCR测试平台_测试计划.docx"
CASE_PATH = OUTPUT_DIR / "银行OCR测试平台_测试用例模板.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK_BLUE = "0B2545"
PALE_BLUE = "E8EEF5"
PALE_GRAY = "F2F4F7"
MID_GRAY = "6B7280"
LIGHT_BORDER = "CDD5DF"
WHITE = "FFFFFF"
BLACK = "111827"
CAUTION_FILL = "FFF8E6"
CAUTION_TEXT = "7A5A00"


def set_run_font(
    run,
    *,
    size: float = 11,
    bold: bool | None = None,
    color: str = BLACK,
    italic: bool | None = None,
) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.shadow = False
    run.font.outline = False
    run.font.emboss = False
    run.font.imprint = False
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], *, indent_dxa: int = 120) -> None:
    total_width = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_width))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for cell, width in zip(row.cells, widths_dxa):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_borders(table, color: str = LIGHT_BORDER, size: str = "4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def style_cell_text(
    cell,
    *,
    bold: bool = False,
    color: str = BLACK,
    size: float = 9.5,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.1
        for run in paragraph.runs:
            set_run_font(run, size=size, bold=bold, color=color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MID_GRAY)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, size=9, color=MID_GRAY)


def add_custom_numbering(document: Document) -> tuple[int, int]:
    numbering = document.part.numbering_part.element
    existing_abstract = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    existing_num = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
    ]
    next_abstract = max(existing_abstract, default=-1) + 1
    next_num = max(existing_num, default=0) + 1

    def make_definition(abstract_id: int, num_id: int, fmt: str, text: str):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)

        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        level.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        level.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), "\uf0b7" if fmt == "bullet" else text)
        level.append(lvl_text)
        justification = OxmlElement("w:lvlJc")
        justification.set(qn("w:val"), "left")
        level.append(justification)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        p_pr.append(ind)
        level.append(p_pr)
        if fmt == "bullet":
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), "Symbol")
            fonts.set(qn("w:hAnsi"), "Symbol")
            fonts.set(qn("w:hint"), "default")
            r_pr.append(fonts)
            level.append(r_pr)
        abstract.append(level)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_num_id = OxmlElement("w:abstractNumId")
        abstract_num_id.set(qn("w:val"), str(abstract_id))
        num.append(abstract_num_id)
        return abstract, num

    bullet_abstract, bullet_num = make_definition(next_abstract, next_num, "bullet", "•")
    decimal_abstract, decimal_num = make_definition(next_abstract + 1, next_num + 1, "decimal", "%1.")
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(bullet_abstract)
        numbering.append(decimal_abstract)
    else:
        insertion_index = list(numbering).index(first_num)
        numbering.insert(insertion_index, bullet_abstract)
        numbering.insert(insertion_index + 1, decimal_abstract)
    numbering.append(bullet_num)
    numbering.append(decimal_num)
    return next_num, next_num + 1


def add_numbered_paragraph(document: Document, text: str, num_id: int) -> None:
    """Add a real Word list paragraph linked to a custom numbering definition."""
    paragraph = document.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_element)
    p_pr.append(num_pr)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run)


def configure_document(document: Document, *, landscape: bool = False, running_label: str) -> tuple[int, int]:
    section = document.sections[0]
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
    else:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.shadow = False
        style.font.outline = False
        style.font.emboss = False
        style.font.imprint = False
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    header_run = header_p.add_run(running_label)
    set_run_font(header_run, size=9, color=MID_GRAY, bold=True)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    add_page_number(footer_p)

    return add_custom_numbering(document)


def add_title_block(document: Document, title: str, subtitle: str, status: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title)
    set_run_font(run, size=26, bold=True, color=INK_BLUE)

    sub = document.add_paragraph()
    sub.paragraph_format.space_after = Pt(16)
    run = sub.add_run(subtitle)
    set_run_font(run, size=12, color=MID_GRAY)

    table = document.add_table(rows=4, cols=2)
    rows = [
        ("项目", "Bank OCR Test Platform"),
        ("文档版本", "V1.0"),
        ("编制日期", date.today().isoformat()),
        ("状态", status),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], PALE_GRAY)
        style_cell_text(row.cells[0], bold=True, color=DARK_BLUE, size=10)
        style_cell_text(row.cells[1], size=10)
    set_table_geometry(table, [2700, 6660])
    set_table_borders(table)


def add_callout(document: Document, label: str, text: str, *, caution: bool = False) -> None:
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    label_run = paragraph.add_run(f"{label}：")
    set_run_font(label_run, size=10.5, bold=True, color=CAUTION_TEXT if caution else DARK_BLUE)
    body_run = paragraph.add_run(text)
    set_run_font(body_run, size=10.5, color=BLACK)
    set_cell_shading(cell, CAUTION_FILL if caution else PALE_BLUE)
    set_table_geometry(table, [9360])
    set_table_borders(table, color="E3C66A" if caution else "B8CBE0")


def add_matrix(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_dxa: list[int],
    *,
    font_size: float = 9,
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        set_cell_shading(cell, PALE_BLUE)
        style_cell_text(cell, bold=True, color=INK_BLUE, size=font_size, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])

    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            cell.text = value
            style_cell_text(cell, size=font_size)

    set_table_geometry(table, widths_dxa)
    set_table_borders(table)


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    size = {1: 16, 2: 13, 3: 12}[level]
    color = BLUE if level in (1, 2) else DARK_BLUE
    set_run_font(run, size=size, bold=True, color=color)


def build_test_plan() -> None:
    document = Document()
    bullet_id, number_id = configure_document(
        document,
        running_label="银行 OCR 测试平台 | 测试计划",
    )
    document.core_properties.title = "银行OCR测试平台测试计划"
    document.core_properties.subject = "测试开发实践项目测试计划"
    document.core_properties.author = "项目测试负责人"

    add_title_block(
        document,
        "银行 OCR 测试平台测试计划",
        "面向银行卡与身份证影像审核流程的功能、接口、数据与性能测试",
        "基线草案，可直接执行并按实际结果更新",
    )
    document.add_paragraph()
    add_callout(
        document,
        "计划目标",
        "通过可重复执行的测试活动，验证上传校验、图像质检、OCR、字段解析、业务规则、审核记录和错误处理链路，并形成可用于复盘与面试展示的真实测试证据。",
    )

    add_heading(document, "1. 项目背景与质量目标", 1)
    document.add_paragraph(
        "本项目是基于 FastAPI 的银行影像 OCR 测试平台 Demo，覆盖银行卡和身份证图片从上传到审核记录落库的完整链路。系统支持 mock 与 PaddleOCR 两种识别模式，使用 SQLite 保存审核记录，并提供 pytest、pytest-html、Allure、Locust 和 GitHub Actions 测试基础设施。"
    )
    add_heading(document, "1.1 质量目标", 2)
    for item in (
        "核心审核接口在合法输入下能够返回稳定、可解释的结果，并包含 request_id、review_result、review_reasons、quality、ocr_text 和 fields。",
        "非法文件、损坏图片、空文件和非法 OCR 模式能够被识别，并返回一致的 HTTP 状态码与原因码。",
        "银行卡与身份证字段解析、正反面判断和审核规则覆盖主要正常、异常与边界场景。",
        "审核记录与接口结果一致，敏感号码不会以完整明文写入日志。",
        "回归测试可重复执行；性能测试能够输出响应时间、吞吐量和失败率基线。",
    ):
        add_numbered_paragraph(document, item, bullet_id)

    add_heading(document, "2. 测试范围", 1)
    scope_rows = [
        ["接口", "POST /bank-card/review", "上传校验、质检、OCR、解析、规则、审计和响应契约"],
        ["接口", "POST /id-card/review", "正反面判断、字段解析、质量原因和审核结果"],
        ["接口", "GET /review-records", "按 doc_type、review_result 查询与排序"],
        ["接口", "GET /review-records/{request_id}", "存在与不存在 request_id 的查询行为"],
        ["模块", "quality_check.py", "模糊、过暗、过亮、反光和原因码"],
        ["模块", "ocr_service.py", "mock/paddle 模式、版本输出归一化和异常"],
        ["模块", "field_parser.py", "银行卡号、有效期和持卡人姓名"],
        ["模块", "id_card_parser.py", "身份证正反面及对应字段"],
        ["模块", "rule_check.py", "pass/review/reject 及原因码优先级"],
        ["数据与安全", "SQLite / logging", "记录一致性、过滤、迁移和敏感信息脱敏"],
        ["非功能", "Locust / CI", "性能基线、自动化回归与持续集成"],
    ]
    add_matrix(document, ["类别", "对象", "主要验证内容"], scope_rows, [1500, 2700, 5160], font_size=9.2)

    add_heading(document, "2.1 暂不纳入本轮范围", 2)
    for item in (
        "真实银行生产环境、真实客户证件、支付交易和外部银行系统集成。",
        "移动端 App 原生兼容性、Charles 抓包、弱网和多机型适配。",
        "生产级鉴权、权限、加密、灾备、分布式部署和监管合规审计。",
        "PaddleOCR 模型训练、模型精调及生产级准确率承诺。",
    ):
        add_numbered_paragraph(document, item, bullet_id)

    add_heading(document, "3. 测试策略与方法", 1)
    strategy_rows = [
        ["单元测试", "函数级验证", "解析、质检、规则、脱敏", "pytest"],
        ["接口测试", "状态码、字段、原因码、请求契约", "4 个核心查询/审核接口", "TestClient / pytest"],
        ["集成测试", "端到端模块串联", "上传→质检→OCR→解析→规则→SQLite", "mock 与临时数据库"],
        ["异常测试", "错误输入和服务配置", "空文件、伪文件、损坏图、非法模式", "pytest"],
        ["边界值测试", "临界长度和阈值", "卡号 15/16/19/20 位、月份 00/01/12/13", "参数化测试"],
        ["数据驱动测试", "多质量类型样本", "normal/blur/dark/bright/glare/rotate/occlusion", "pytest 参数化"],
        ["数据库测试", "记录和查询一致性", "唯一性、字段值、筛选、迁移", "SQLite / SQL"],
        ["安全性基础测试", "敏感信息保护", "银行卡号、身份证号日志脱敏", "caplog / 日志检查"],
        ["性能测试", "并发和稳定性基线", "1/10/20/50 用户，失败率、P95、RPS", "Locust"],
        ["回归测试", "变更后全量验证", "137 条现有用例及新增用例", "pytest / GitHub Actions"],
    ]
    add_matrix(document, ["类型", "目标", "重点", "工具"], strategy_rows, [1500, 2100, 3900, 1860], font_size=8.7)

    add_heading(document, "4. 测试环境", 1)
    environment_rows = [
        ["本地操作系统", "Windows，项目工作目录 J:\\job\\bank-ocr-test-platform"],
        ["CI 环境", "GitHub Actions / ubuntu-latest / Python 3.10"],
        ["Python 环境", "Conda 环境 bank；默认系统 Python 不作为测试基线"],
        ["Web 框架", "FastAPI + Uvicorn"],
        ["OCR 模式", "OCR_MODE=mock（回归基线）；OCR_MODE=paddle（真实识别验证）"],
        ["数据存储", "SQLite；默认 reports/review_records.db，测试使用临时数据库"],
        ["测试工具", "pytest、pytest-html、Allure、Locust、FastAPI TestClient"],
        ["报告位置", "reports/test-report.html、Allure 结果及本计划约定的测试总结"],
    ]
    add_matrix(document, ["项目", "配置"], environment_rows, [2200, 7160], font_size=9.5)

    add_heading(document, "5. 测试数据", 1)
    document.add_paragraph(
        "只使用仓库中的合成或明确标注为测试用途的数据，不使用真实银行卡、身份证或客户信息。执行前应记录样本路径、质量类型和期望字段。"
    )
    data_rows = [
        ["银行卡处理集", "data/processed/bank_card/", "normal、blur、dark、bright、glare、rotate、occlusion；各 100 张"],
        ["身份证正面集", "data/processed/id_card/front/", "7 种质量类型；各 100 张"],
        ["身份证反面集", "data/processed/id_card/back/", "7 种质量类型；各 100 张"],
        ["标注文件", "data/annotations/labels.json", "样本路径、文档类型、质量类型和期望字段"],
        ["临时上传", "reports/tmp_uploads/", "接口执行后应清理，不应长期保留敏感影像"],
    ]
    add_matrix(document, ["数据集", "路径", "用途/规模"], data_rows, [1900, 3000, 4460], font_size=9)

    add_heading(document, "6. 准入、暂停与退出标准", 1)
    add_heading(document, "6.1 准入标准", 2)
    for item in (
        "需求和接口字段已确认，测试环境能够启动，依赖安装完成。",
        "核心合成图片可读取，mock 模式可稳定运行，测试数据库可写。",
        "测试用例已评审或完成自检，预期结果明确。",
    ):
        add_numbered_paragraph(document, item, bullet_id)
    add_heading(document, "6.2 暂停条件", 2)
    for item in (
        "核心接口无法启动或超过一半核心用例被同一环境问题阻塞。",
        "测试数据损坏、OCR 模型不可用或数据库持续不可写。",
        "出现可能泄露真实敏感信息的情况，应立即停止并清理。",
    ):
        add_numbered_paragraph(document, item, bullet_id)
    add_heading(document, "6.3 退出标准", 2)
    for item in (
        "P0/P1 核心用例 100% 执行，核心功能无未关闭的致命或严重缺陷。",
        "全量自动化回归通过；若有失败，已确认原因并记录处置结论。",
        "性能基线已执行，失败率与关键响应指标已记录。",
        "测试总结包含用例统计、缺陷统计、遗留风险和是否建议通过的结论。",
    ):
        add_numbered_paragraph(document, item, bullet_id)

    add_heading(document, "7. 缺陷管理", 1)
    defect_rows = [
        ["致命 S1", "服务不可用、数据严重损坏、敏感信息大规模泄露", "立即修复，阻断测试/发布"],
        ["严重 S2", "核心审核结果错误、记录无法保存、主要接口持续 5xx", "优先修复，修复后专项回归"],
        ["一般 S3", "非核心分支错误、原因码不一致、部分数据解析异常", "纳入当前或下一迭代"],
        ["轻微 S4", "提示文案、日志格式、轻微 UI 或易用性问题", "评估后优化"],
    ]
    add_matrix(document, ["级别", "判定示例", "处理要求"], defect_rows, [1600, 5100, 2660], font_size=9)
    document.add_paragraph(
        "缺陷生命周期建议使用：新建 → 已确认 → 修复中 → 待验证 → 已关闭；若验证失败则重新打开。每个缺陷至少记录环境、前置条件、步骤、实际结果、期望结果、证据和关联用例。"
    )

    add_heading(document, "8. 执行安排", 1)
    schedule_rows = [
        ["第 1 天", "需求与代码走查、风险分析、测试计划", "测试计划 V1.0"],
        ["第 2 天", "核心功能与接口用例设计", "用例基线"],
        ["第 3 天", "银行卡、身份证、异常和边界用例执行", "执行记录、缺陷"],
        ["第 4 天", "数据库、日志脱敏和前端验证", "SQL/日志证据"],
        ["第 5 天", "补充自动化测试并执行全量回归", "pytest-html / Allure"],
        ["第 6 天", "Locust 性能基线与稳定性观察", "性能数据"],
        ["第 7 天", "缺陷复测、总结与面试复盘", "测试报告、项目讲述稿"],
    ]
    add_matrix(document, ["阶段", "活动", "交付物"], schedule_rows, [1500, 5100, 2760], font_size=9)

    add_heading(document, "9. 度量与测试报告", 1)
    metric_rows = [
        ["用例执行率", "已执行用例数 / 计划用例数", "目标 100%"],
        ["用例通过率", "通过用例数 / 已执行用例数", "用于判断当前质量，不单独作为放行依据"],
        ["自动化通过率", "自动化通过数 / 自动化执行数", "核心回归目标 100%"],
        ["缺陷分布", "按严重级别、模块、状态统计", "识别高风险模块"],
        ["缺陷重开率", "重新打开缺陷数 / 已验证缺陷数", "评估修复质量"],
        ["性能指标", "平均、P50、P90、P95、RPS、失败率", "不同并发档位分别记录"],
    ]
    add_matrix(document, ["指标", "口径", "说明"], metric_rows, [1900, 3500, 3960], font_size=9)
    add_callout(
        document,
        "报告要求",
        "自动化 HTML/Allure 报告只证明自动化执行结果；最终测试报告还必须说明测试范围、环境、人工用例、缺陷、性能数据、遗留风险和质量结论。",
        caution=True,
    )

    add_heading(document, "10. 风险与应对", 1)
    risk_rows = [
        ["PaddleOCR 首次加载慢或依赖不稳定", "性能波动、环境阻塞", "mock 作为回归基线；paddle 独立记录环境和预热条件"],
        ["合成数据与真实影像差异", "准确率结论代表性有限", "明确 Demo 边界，不宣称生产准确率"],
        ["固定阈值依赖数据分布", "不同图片误判", "记录阈值附近样本并进行边界分析"],
        ["SQLite 并发写入限制", "性能测试出现锁库", "监控 database is locked 并单独统计"],
        ["测试报告数据被手工编造", "面试无法解释细节", "所有统计必须来自真实执行记录和报告文件"],
    ]
    add_matrix(document, ["风险", "影响", "应对"], risk_rows, [2600, 2700, 4060], font_size=9)

    document.add_page_break()
    add_heading(document, "11. 交付物清单", 1)
    deliverable_rows = [
        ["测试计划", "本文件，随范围、环境和风险变化维护版本"],
        ["测试用例", "测试用例模板及实际执行记录"],
        ["自动化测试", "pytest 用例、pytest-html 报告和可选 Allure 报告"],
        ["缺陷记录", "缺陷编号、严重级别、复现步骤、证据和验证结果"],
        ["性能结果", "Locust 并发档位、响应时间、吞吐量和失败率"],
        ["测试总结", "范围、环境、统计、遗留风险和质量结论"],
    ]
    add_matrix(document, ["交付物", "内容"], deliverable_rows, [2200, 7160], font_size=9.2)

    add_heading(document, "12. 评审与变更记录", 1)
    add_matrix(
        document,
        ["版本", "日期", "变更说明", "编制/评审"],
        [["V1.0", date.today().isoformat(), "建立银行 OCR 测试平台测试计划基线", "【填写】"]],
        [1200, 1800, 4560, 1800],
        font_size=9.5,
    )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(PLAN_PATH)


def add_case_register_table(document: Document, rows: list[list[str]]) -> None:
    headers = ["用例ID", "模块", "用例标题", "类型", "优先级", "前置条件", "步骤/数据", "预期结果", "状态"]
    widths = [950, 1200, 1900, 1000, 800, 1550, 2500, 2300, 760]
    table = document.add_table(rows=1, cols=len(headers))
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        set_cell_shading(cell, PALE_BLUE)
        style_cell_text(cell, bold=True, color=INK_BLUE, size=8.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])

    for row_values in rows:
        cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, row_values)):
            cell.text = value
            align = WD_ALIGN_PARAGRAPH.CENTER if index in (0, 1, 3, 4, 8) else WD_ALIGN_PARAGRAPH.LEFT
            style_cell_text(cell, size=7.6, align=align)

    set_table_geometry(table, widths)
    set_table_borders(table)


def build_test_case_template() -> None:
    document = Document()
    bullet_id, number_id = configure_document(
        document,
        landscape=True,
        running_label="银行 OCR 测试平台 | 测试用例模板",
    )
    document.core_properties.title = "银行OCR测试平台测试用例模板"
    document.core_properties.subject = "测试用例设计与执行记录模板"
    document.core_properties.author = "项目测试负责人"

    add_title_block(
        document,
        "银行 OCR 测试平台测试用例模板",
        "包含字段规范、空白模板、代表性示例和执行统计口径",
        "模板基线；【填写】内容应在实际执行后更新",
    )
    add_callout(
        document,
        "使用方法",
        "先复制空白用例行，再补齐前置条件、步骤、数据和预期结果；执行后填写实际结果、状态、缺陷编号、执行人和日期。示例用例用于帮助起步，不能代替真实执行。",
    )

    add_heading(document, "1. 用例字段说明", 1)
    fields_rows = [
        ["用例ID", "唯一编号，建议 模块缩写-三位序号", "BC-API-001"],
        ["模块", "所属功能或代码模块", "银行卡接口"],
        ["用例标题", "一句话说明条件和期望", "上传正常 PNG 应审核通过"],
        ["类型", "功能/接口/异常/边界/数据库/安全/性能/回归", "边界"],
        ["优先级", "P0 核心、P1 重要、P2 一般、P3 优化", "P0"],
        ["前置条件", "环境、OCR 模式、数据库和数据准备", "OCR_MODE=mock"],
        ["测试步骤", "可复现的编号步骤", "上传图片→提交→读取响应"],
        ["测试数据", "文件路径、参数或构造数据", "normal/bank_card_0001.png"],
        ["预期结果", "状态码、字段、原因码和数据库结果", "HTTP 200，review_result=pass"],
        ["实际结果", "执行后观察到的真实结果", "【填写】"],
        ["状态", "未执行/通过/失败/阻塞/不适用", "未执行"],
        ["缺陷编号", "失败时关联缺陷；无缺陷留空", "BUG-001"],
        ["执行信息", "执行人、日期、构建/提交号", "【填写】"],
    ]
    add_matrix(document, ["字段", "填写规则", "示例"], fields_rows, [1800, 6500, 4660], font_size=9)

    add_heading(document, "2. 单条测试用例空白模板", 1)
    blank_rows = [
        ["用例ID", "【填写】", "模块", "【填写】"],
        ["用例标题", "【填写】", "优先级", "P0 / P1 / P2 / P3"],
        ["测试类型", "功能 / 接口 / 异常 / 边界 / 数据库 / 安全 / 性能 / 回归", "关联需求", "【填写】"],
        ["前置条件", "【填写】", "测试环境", "【填写】"],
        ["测试步骤", "1. 【填写】\n2. 【填写】\n3. 【填写】", "测试数据", "【填写】"],
        ["预期结果", "【填写】", "实际结果", "【执行后填写】"],
        ["执行状态", "未执行 / 通过 / 失败 / 阻塞 / 不适用", "缺陷编号", "【如有则填写】"],
        ["执行人/日期", "【填写】", "构建/提交号", "【填写】"],
        ["备注/证据", "【截图、日志、request_id、报告链接等】", "", ""],
    ]
    table = document.add_table(rows=0, cols=4)
    for values in blank_rows:
        cells = table.add_row().cells
        if values[2] == "" and values[3] == "":
            cells[1].merge(cells[3])
        for index, value in enumerate(values):
            if index >= len(cells):
                break
            cells[index].text = value
        for index, cell in enumerate(cells):
            if index in (0, 2):
                set_cell_shading(cell, PALE_GRAY)
                style_cell_text(cell, bold=True, color=DARK_BLUE, size=9)
            else:
                style_cell_text(cell, size=9)
    set_table_geometry(table, [1700, 4780, 1700, 4780])
    set_table_borders(table)

    add_heading(document, "3. 代表性测试用例示例", 1)
    document.add_paragraph(
        "以下示例覆盖当前项目的主要风险。执行时应把“步骤/数据”拆成可复现步骤，并把 response、request_id、数据库记录或日志作为证据。"
    )
    example_rows = [
        ["BC-API-001", "银行卡接口", "上传正常 PNG 返回通过", "接口", "P0", "服务启动；mock 模式", "上传 data/processed/bank_card/normal/bank_card_0001.png", "HTTP 200；pass；字段完整；记录可查询", "未执行"],
        ["BC-API-002", "银行卡接口", "缺少 file 字段返回 422", "异常", "P0", "服务启动", "POST /bank-card/review，不传 multipart file", "HTTP 422；含 request_id 和 invalid_request", "未执行"],
        ["BC-API-003", "银行卡接口", "上传空文件被拒绝", "异常", "P0", "服务启动", "上传文件名 empty.png，内容 0 字节", "HTTP 400；unreadable_image；写入错误记录", "未执行"],
        ["BC-API-004", "银行卡接口", "上传 TXT 扩展名被拒绝", "异常", "P1", "服务启动", "上传 sample.txt", "HTTP 400；invalid_file_type", "未执行"],
        ["BC-RULE-001", "银行卡规则", "卡号 16 位为有效边界", "边界", "P0", "调用规则函数", "card_number=16 位数字；其他字段与质量正常", "review_result=pass", "未执行"],
        ["BC-RULE-002", "银行卡规则", "卡号 20 位被拒绝", "边界", "P0", "调用规则函数", "card_number=20 位数字", "review_result=reject；invalid_card_number", "未执行"],
        ["BC-RULE-003", "银行卡规则", "有效期月份 13 需要复核", "边界", "P1", "调用规则函数", "valid_date=13/30", "review；invalid_valid_date", "未执行"],
        ["BC-QUALITY-001", "图片质检", "模糊图片进入人工复核", "功能", "P0", "样本存在", "上传 blur/bank_card_0001.png", "quality_result=review；image_blur", "未执行"],
        ["BC-QUALITY-002", "图片质检", "反光区域超过阈值进入复核", "边界", "P1", "构造阈值附近样本", "分别验证最大连通区域比例 0.0049/0.0051", "前者不因反光复核；后者 glare_detected", "未执行"],
        ["ID-API-001", "身份证接口", "正面字段完整返回通过", "接口", "P0", "mock OCR 为正面文本", "上传可读正面图片", "side=front；6 个必填字段完整；pass", "未执行"],
        ["ID-API-002", "身份证接口", "反面字段完整返回通过", "接口", "P0", "mock OCR 为反面文本", "上传可读反面图片", "side=back；签发机关和有效期限完整；pass", "未执行"],
        ["ID-PARSER-001", "身份证解析", "无正反面关键词返回 unknown", "异常", "P1", "调用解析函数", "OCR 文本为 TEST DATA", "side=unknown；fields 为空", "未执行"],
        ["ID-PARSER-002", "身份证解析", "正反面得分相同时选择正面", "边界", "P2", "调用 side 判断", "构造 front_score=back_score>0", "当前实现返回 front；记录为已知规则", "未执行"],
        ["REC-DB-001", "审核记录", "成功请求写入 SQLite", "数据库", "P0", "使用临时数据库", "提交正常银行卡并按 request_id 查询", "接口与数据库 review_result、reasons、fields 一致", "未执行"],
        ["REC-DB-002", "审核记录", "筛选结果按 id 倒序", "数据库", "P1", "预置多条记录", "GET /review-records?doc_type=bank_card&review_result=review", "仅返回匹配记录，按 id DESC", "未执行"],
        ["SEC-LOG-001", "日志脱敏", "银行卡号不以完整明文出现", "安全", "P0", "开启日志捕获", "提交含 16 位卡号的审核请求", "日志仅保留前 6 后 4，中间为星号", "未执行"],
        ["SEC-LOG-002", "日志脱敏", "身份证号不以完整明文出现", "安全", "P0", "开启日志捕获", "解析并记录身份证字段", "出生日期段被星号遮盖", "未执行"],
        ["OCR-001", "OCR 服务", "默认 mock 不加载 PaddleOCR", "回归", "P1", "未设置 OCR_MODE", "调用 recognize_text", "返回固定文本；不初始化 Paddle 引擎", "未执行"],
        ["OCR-002", "OCR 服务", "非法 OCR_MODE 返回服务端错误", "异常", "P0", "OCR_MODE=invalid", "提交审核请求", "HTTP 500；invalid_ocr_mode；含 request_id", "未执行"],
        ["PERF-001", "性能", "10 用户 mock 模式基线", "性能", "P1", "服务启动；mock；数据准备", "Locust 10 用户运行 5 分钟", "记录平均/P95/RPS/失败率；无未解释失败", "未执行"],
    ]
    add_case_register_table(document, example_rows)

    add_heading(document, "4. 执行记录模板", 1)
    execution_headers = ["用例ID", "构建/提交号", "执行环境", "实际结果摘要", "状态", "缺陷编号", "执行人", "执行日期", "证据位置"]
    execution_rows = [
        ["【填写】", "【填写】", "本地 / CI / Paddle", "【填写】", "未执行", "", "【填写】", "【填写】", "【截图/日志/报告路径】"],
        ["【填写】", "【填写】", "本地 / CI / Paddle", "【填写】", "未执行", "", "【填写】", "【填写】", "【截图/日志/报告路径】"],
        ["【填写】", "【填写】", "本地 / CI / Paddle", "【填写】", "未执行", "", "【填写】", "【填写】", "【截图/日志/报告路径】"],
    ]
    add_matrix(document, execution_headers, execution_rows, [1100, 1500, 1500, 2700, 1000, 1300, 1100, 1300, 2460], font_size=8)

    add_heading(document, "5. 缺陷记录模板", 1)
    defect_headers = ["缺陷ID", "关联用例", "标题", "严重级别", "环境", "复现步骤", "实际/期望", "状态", "验证结果"]
    defect_rows = [
        ["BUG-【序号】", "【用例ID】", "【条件】下【模块】出现【错误结果】", "S1/S2/S3/S4", "【填写】", "1. 【填写】\n2. 【填写】", "实际：【填写】\n期望：【填写】", "新建", "【填写】"],
        ["BUG-【序号】", "【用例ID】", "【填写】", "S1/S2/S3/S4", "【填写】", "1. 【填写】", "实际：【填写】\n期望：【填写】", "新建", "【填写】"],
    ]
    add_matrix(document, defect_headers, defect_rows, [1100, 1200, 2100, 1100, 1300, 2400, 2200, 900, 1660], font_size=8)

    add_heading(document, "6. 测试执行统计模板", 1)
    statistic_rows = [
        ["计划用例数", "【填写】", "执行率", "已执行 / 计划", "【填写】"],
        ["已执行", "【填写】", "通过率", "通过 / 已执行", "【填写】"],
        ["通过", "【填写】", "失败率", "失败 / 已执行", "【填写】"],
        ["失败", "【填写】", "阻塞率", "阻塞 / 已执行", "【填写】"],
        ["阻塞", "【填写】", "自动化通过率", "自动化通过 / 自动化执行", "【填写】"],
        ["缺陷总数", "【填写】", "未关闭 S1/S2", "按缺陷记录统计", "【填写】"],
    ]
    add_matrix(document, ["指标", "结果", "比率指标", "计算口径", "结果"], statistic_rows, [1800, 1400, 2200, 5400, 2160], font_size=9)
    add_callout(
        document,
        "面试诚信提示",
        "只有亲自执行并保留证据的用例，才可以在面试中表述为“我做过”。可说明这是个人实践项目，但不要把预置模板或已有测试冒充为商业项目经历。",
        caution=True,
    )

    add_heading(document, "7. 模板变更记录", 1)
    add_matrix(
        document,
        ["版本", "日期", "变更说明", "维护人"],
        [["V1.0", date.today().isoformat(), "建立用例、执行、缺陷和统计模板；加入代表性示例", "【填写】"]],
        [1500, 1800, 6500, 3160],
        font_size=9,
    )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(CASE_PATH)


def main() -> None:
    build_test_plan()
    build_test_case_template()
    print(PLAN_PATH)
    print(CASE_PATH)


if __name__ == "__main__":
    main()
